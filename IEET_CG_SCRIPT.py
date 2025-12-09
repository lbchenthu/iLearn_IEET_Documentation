import math
import random
import os
import os.path
import re
import argparse

import pandas as pd
import numpy as np
from docx import Document


# ================== 成績解析工具 ================== #

def parse_score(val):
    """
    把單一格轉成分數 (float) 或 None：
    - 空白、NaN、'未繳', '未交', '缺交', '未考', '未批改' => None
    - 其他能轉成數字的 => float 分數
    """
    if pd.isna(val):
        return None
    if isinstance(val, (int, float, np.integer, np.floating)):
        return float(val)

    s = str(val).strip()
    if s in ["", "未繳", "未交", "缺交", "未考", "未批改"]:
        return None

    try:
        return float(s)
    except ValueError:
        return None


def is_assignment(col_name: str) -> bool:
    """
    判斷是否視為「作業/專案」欄位。
    """
    if not isinstance(col_name, str):
        return False

    name = col_name.strip().lower()
    ASSIGNMENT_KEYWORDS = [
        "作業",
        "hw",
        "lab",
        "報告",
        "project",
        "專題",
        "專案",
        "上傳區",
        "成績(",   # 避掉「總成績」，總成績另外排除
    ]
    return any(k in name for k in ASSIGNMENT_KEYWORDS)


def should_select_column(col_name: str, scores_raw) -> bool:
    """
    決定此「評分項目」是否要統計：
    (a). 要挑選: 包含 "考"的項目
    (b). 要挑選: 作業 (排除過半數 未評分/未繳交 的作業)
    (c). 不要挑選: 包含"點名"的項目
    (d). 不要挑選: 包含"總成績"的項目
    """
    if not isinstance(col_name, str):
        return False

    name = col_name.strip()

    # (c) 不要: 點名
    if "點名" in name:
        return False

    # (d) 不要: 各種總成績
    BAN_KEYWORDS = ["總成績", "學期總成績", "期末總成績", "原始總成績"]
    if any(b in name for b in BAN_KEYWORDS):
        return False

    # (a) 要: 考試
    if "考" in name:
        return True

    # (b) 作業
    if is_assignment(name):
        parsed = [parse_score(v) for v in scores_raw]
        expected = len(parsed)
        submitted = sum(1 for s in parsed if s is not None)
        missing = expected - submitted
        # 過半數未繳/未評分就不選
        if missing > expected / 2:
            return False
        return True

    return False


# ================== 作業序號 / 類型解析 ================== #

def parse_item_meta(col_name: str):
    """
    解析評分項目的「作業種類」，只分三種：
    - 期中考
    - 期末考
    - 作業（只要名稱不是期中考、期末考，一律歸類為作業）
    """
    name = str(col_name).strip()

    if "期中" in name and "考" in name:
        i_type = "期中考"
    elif "期末" in name and "考" in name:
        i_type = "期末考"
    else:
        # 其餘一律當作「作業」
        i_type = "作業"

    return {"I_TYPE": i_type}


# ================== 三區段分組（同分同區） ================== #

def split_into_three_segments(valid_students):
    """
    valid_students: list of dict, each dict 至少有
      {
        "id": 學號,
        "name": 姓名,
        "score": 分數 (float)
      }

    規則：
    - 只針對有效成績分段
    - 同分數必須分在同一區段
    - 儘量使三區段人數接近
    - 回傳 [segA, segB, segC]
    """
    if not valid_students:
        return [[], [], []]

    # 依分數由高到低排序
    sorted_students = sorted(valid_students, key=lambda x: x["score"], reverse=True)

    # 先照「同分數」分群
    groups = []
    current_group = [sorted_students[0]]

    for stu in sorted_students[1:]:
        if math.isclose(stu["score"], current_group[0]["score"]):
            current_group.append(stu)
        else:
            groups.append(current_group)
            current_group = [stu]
    groups.append(current_group)

    total = len(sorted_students)
    segments = [[], [], []]
    seg_counts = [0, 0, 0]

    seg_idx = 0
    remaining_students = total
    remaining_segments = 3

    for g in groups:
        g_size = len(g)

        # 最後一區：全部塞進去
        if seg_idx == 2:
            segments[seg_idx].extend(g)
            seg_counts[seg_idx] += g_size
            remaining_students -= g_size
            continue

        target = remaining_students / remaining_segments

        if seg_counts[seg_idx] > 0 and seg_counts[seg_idx] + g_size > target and seg_idx < 2:
            seg_idx += 1
            remaining_segments -= 1
            target = remaining_students / remaining_segments

        segments[seg_idx].extend(g)
        seg_counts[seg_idx] += g_size
        remaining_students -= g_size

    return segments  # [A, B, C]


# ================== 課程資訊解析 ================== #

def _extract_neighbor_or_suffix(df_raw, r, c):
    """
    嘗試從某個 cell 取出「值」：
    - 若該格本身有 "標籤：內容" 形式，就取 "：" 後面的部份
    - 否則回傳右邊那一格 (同一列下一欄)
    """
    val = df_raw.iat[r, c]
    if isinstance(val, str):
        s = val.strip()
        for sep in ["：", ":"]:
            if sep in s:
                return s.split(sep, 1)[1].strip()

    if c + 1 < df_raw.shape[1]:
        right_val = df_raw.iat[r, c + 1]
        if isinstance(right_val, str):
            return right_val.strip()
        return right_val
    return None


def clean_course_name(name: str) -> str:
    """
    課程名稱: {{C_NAME}}，要去除編號與括號。

    規則：
    - 去掉所有括號(...) / （...）
    - 去掉前面的「數字+符號」，例如 "001. ", "1) ", "01-"
    """
    if not isinstance(name, str):
        return str(name)

    s = name

    # 去掉括號內容 (含中英文括號)
    s = re.sub(r"[\(（][^\)）]*[\)）]", "", s)

    # 去掉前置數字編號 e.g. "001. ", "1) ", "01-"
    s = re.sub(r"^\s*\d+\s*[\.\-)]\s*", "", s)

    # 再 strip 一次
    s = s.strip()
    return s


def parse_course_info(df_raw, header_row_idx):
    """
    從 header_row_idx 之前的列中解析課程資訊，填到占位符 mapping。

    佔位符：
      {{C_NAME}}, {{C_TEACHER}}, {{DEPT}}, {{C_COURSE_ID}},
      {{C_YEAR}}, {{C_SEM}}, {{C_POINT}}, {{C_ADMIN_ID}}, {{C_TA}}
    （C_CLASS 由學生資料決定）
    """
    course_mapping = {
        "{{C_NAME}}": "",
        "{{C_TEACHER}}": "",
        "{{DEPT}}": "",
        "{{C_COURSE_ID}}": "",
        "{{C_YEAR}}": "",
        "{{C_SEM}}": "",
        "{{C_POINT}}": "",
        "{{C_ADMIN_ID}}": "",
        "{{C_TA}}": "",
    }

    for r in range(header_row_idx):
        for c in range(df_raw.shape[1]):
            val = df_raw.iat[r, c]
            if not isinstance(val, str):
                continue
            text = val.strip()

            # 課程名稱 / 科目名稱
            if any(k in text for k in ["課程名稱", "科目名稱"]) and not course_mapping["{{C_NAME}}"]:
                v = _extract_neighbor_or_suffix(df_raw, r, c)
                if v is not None:
                    course_mapping["{{C_NAME}}"] = clean_course_name(str(v))

            # 授課老師 / 授課教師
            if any(k in text for k in ["授課老師", "授課教師"]) and not course_mapping["{{C_TEACHER}}"]:
                v = _extract_neighbor_or_suffix(df_raw, r, c)
                if v is not None:
                    course_mapping["{{C_TEACHER}}"] = str(v)

            # 院系 / 開課系所 / 科系
            if any(k in text for k in ["院系", "開課系所", "科系", "系所"]) and not course_mapping["{{DEPT}}"]:
                v = _extract_neighbor_or_suffix(df_raw, r, c)
                if v is not None:
                    course_mapping["{{DEPT}}"] = str(v)

            # 課程代碼 / 課程代號
            if any(k in text for k in ["課程代碼", "課程代號"]) and not course_mapping["{{C_COURSE_ID}}"]:
                v = _extract_neighbor_or_suffix(df_raw, r, c)
                if v is not None:
                    course_mapping["{{C_COURSE_ID}}"] = str(v)

            # 教務科目代碼 / 科目代號
            if any(k in text for k in ["教務科目代碼", "科目代號"]) and not course_mapping["{{C_ADMIN_ID}}"]:
                v = _extract_neighbor_or_suffix(df_raw, r, c)
                if v is not None:
                    course_mapping["{{C_ADMIN_ID}}"] = str(v)

            # 學分
            if "學分" in text and not course_mapping["{{C_POINT}}"]:
                v = _extract_neighbor_or_suffix(df_raw, r, c)
                if isinstance(v, (int, float, np.integer, np.floating)):
                    course_mapping["{{C_POINT}}"] = str(v)
                elif isinstance(v, str):
                    m = re.search(r"(\d+(\.\d+)?)", v)
                    if m:
                        course_mapping["{{C_POINT}}"] = m.group(1)

            # 學年 + 學期（抓一個四位數）
            if "學年" in text and "學期" in text and (not course_mapping["{{C_YEAR}}"] or not course_mapping["{{C_SEM}}"]):
                candidate = text + " "
                right = df_raw.iat[r, c + 1] if c + 1 < df_raw.shape[1] else ""
                if isinstance(right, str):
                    candidate += right
                m = re.search(r"(\d{4})", candidate)
                if m:
                    code = m.group(1)  # e.g. 1121
                    course_mapping["{{C_YEAR}}"] = code[:3]
                    course_mapping["{{C_SEM}}"] = code[3]

            # 助教
            if "助教" in text and not course_mapping["{{C_TA}}"]:
                v = _extract_neighbor_or_suffix(df_raw, r, c)
                if v is not None:
                    course_mapping["{{C_TA}}"] = str(v)

    return course_mapping


# ================== Word 範本填入（保留字型 + 勾選 checkbox） ================== #

def replace_placeholders_in_runs(runs, mapping):
    """
    在不破壞 run 格式的情況下，處理「跨 run 的 placeholder」。

    作法：
    - 對每個 key：
      - 把所有 run.text 串成一條字串 S
      - 找到 key 在 S 的位置
      - 用 char index → (run_idx, char_idx) 的對應，把這段改成 value
    注意：這裡假設 placeholder 在 template 中不會被拆成很奇怪的排版。
    """
    for key, val in mapping.items():
        while True:
            S = "".join(r.text for r in runs)
            idx = S.find(key)
            if idx == -1:
                break

            # 建 char index -> (run_idx, char_idx)
            pos2rc = []
            for ri, r in enumerate(runs):
                for ci, ch in enumerate(r.text):
                    pos2rc.append((ri, ci))

            start = idx
            end = idx + len(key)
            if end > len(pos2rc):
                break

            r_start, c_start = pos2rc[start]
            r_end, c_end = pos2rc[end - 1]

            repl = str(val)
            prefix = runs[r_start].text[:c_start]
            runs[r_start].text = prefix + repl

            # 清空 placeholder 後續 run 的文字
            for ri in range(r_start + 1, r_end + 1):
                runs[ri].text = ""


def apply_checkbox_in_runs(runs, i_type):
    """
    根據單一 I_TYPE（期中考 / 期末考 / 作業），
    把「評量類別」那一行的 checkbox 處理成只勾一個。

    作法：
    - 只在該段文字包含「評量類別」時動作
    - 先把所有 '■' 清回 '□'
    - 再依 I_TYPE 把對應 '□ XX' 換成 '■ XX'
    - 把整行文字寫回第一個 run，其他 run 清空
    """
    S = "".join(r.text for r in runs)
    if "評量類別" not in S:
        return

    # 全部清成空框
    S = S.replace("■", "□")

    if i_type == "期中考":
        S = S.replace("□ 期中考", "■ 期中考", 1)
    elif i_type == "期末考":
        S = S.replace("□ 期末考", "■ 期末考", 1)
    else:  # 作業
        S = S.replace("□ 作業", "■ 作業", 1)

    # 寫回 runs：第一個 run 放完整內容，其餘清空
    if runs:
        runs[0].text = S
        for r in runs[1:]:
            r.text = ""


def fill_template(template_path, output_path, mapping, item_meta):
    """
    template_path: 範本 .docx
    output_path: 輸出 .docx
    mapping: 佔位符替換 dict
    item_meta: {"I_ID": ..., "I_TYPE": ...} 用來處理 checkbox
    """
    doc = Document(template_path)

    i_type = item_meta.get("I_TYPE", "")

    # 段落
    for para in doc.paragraphs:
        replace_placeholders_in_runs(para.runs, mapping)
        apply_checkbox_in_runs(para.runs, i_type)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_runs(para.runs, mapping)
                    apply_checkbox_in_runs(para.runs, i_type)

    doc.save(output_path)


def build_mapping_for_item(item_name, stats, course_mapping, item_meta):
    """
    stats:
      expected, submitted, missing, full_avg, segments
    course_mapping:
      各 {{C_xxx}} 相關
    item_meta:
      {"I_ID": ..., "I_TYPE": ...}
    """
    m = {
        "{{ITEM_NAME}}": item_name,
        "{{EXPECTED}}": stats["expected"],
        "{{SUBMITTED}}": stats["submitted"],
        "{{MISSING}}": stats["missing"],
        "{{FULL_AVG}}": f"{stats['full_avg']:.2f}",

        "{{I_ID}}": item_meta.get("I_ID", ""),
        "{{I_TYPE}}": item_meta.get("I_TYPE", ""),
    }

    for label in ["A", "B", "C"]:
        seg = stats["segments"][label]

        if seg["count"] > 0:
            m[f"{{{{{label}_MAX}}}}"] = f"{seg['max']:.2f}"
            m[f"{{{{{label}_MIN}}}}"] = f"{seg['min']:.2f}"
            m[f"{{{{{label}_COUNT}}}}"] = seg["count"]
            m[f"{{{{{label}_AVG}}}}"] = f"{seg['avg']:.2f}"
        else:
            m[f"{{{{{label}_MAX}}}}"] = ""
            m[f"{{{{{label}_MIN}}}}"] = ""
            m[f"{{{{{label}_COUNT}}}}"] = 0
            m[f"{{{{{label}_AVG}}}}"] = ""

        # 隨機抽 2 位（不足補空）
        students = seg["students"]
        if seg["count"] == 0:
            samples = []
        else:
            n = min(2, seg["count"])
            samples = random.sample(students, n)

        while len(samples) < 2:
            samples.append(None)

        for i, stu in enumerate(samples, start=1):
            if stu is None:
                m[f"{{{{{label}_STU{i}_ID}}}}"] = ""
                m[f"{{{{{label}_STU{i}_NAME}}}}"] = ""
                m[f"{{{{{label}_STU{i}_SCORE}}}}"] = ""
            else:
                m[f"{{{{{label}_STU{i}_ID}}}}"] = stu["id"]
                m[f"{{{{{label}_STU{i}_NAME}}}}"] = stu["name"]
                m[f"{{{{{label}_STU{i}_SCORE}}}}"] = f"{stu['score']:.2f}"

    # 合併課程資訊占位符
    if course_mapping:
        for k, v in course_mapping.items():
            m.setdefault(k, v)

    return m


# ================== 主邏輯：讀檔 + 統計 + 產出 Word ================== #

def analyze_and_generate(excel_path, template_path, output_dir, sheet_name=None):
    # 不設 header，先拿原始表來找「學號」列 + 上方課程資訊
    if sheet_name is None:
        df_raw = pd.read_excel(excel_path, header=None)
    else:
        df_raw = pd.read_excel(excel_path, sheet_name=sheet_name, header=None)

    # 找出「學號」所在列
    header_row_idx = None
    for i in range(len(df_raw)):
        row_vals = df_raw.iloc[i].tolist()
        if any(str(v).strip() == "學號" for v in row_vals):
            header_row_idx = i
            break
    if header_row_idx is None:
        raise ValueError("找不到「學號」標題列，請確認成績單格式。")

    # 課程資訊
    course_mapping = parse_course_info(df_raw, header_row_idx)

    # 真正的成績 DataFrame：標題列 + 學生列
    header_row = df_raw.iloc[header_row_idx]
    df = df_raw.iloc[header_row_idx + 1:].copy()
    df.columns = header_row

    # 只保留有學號的
    if "學號" not in df.columns or "姓名" not in df.columns:
        raise ValueError("成績資料區找不到「學號」或「姓名」欄。")

    student_df = df[df["學號"].notna() & (df["學號"].astype(str).str.strip() != "")]
    num_students = len(student_df)
    if num_students == 0:
        print("沒有學生資料。")
        return

    # C_CLASS：由學生班級的人數決定（最多數）
    if "班級" in student_df.columns:
        class_series = (
            student_df["班級"]
                .astype(str)
                .str.strip()
                .replace({"nan": ""})
        )
        vc = class_series[class_series != ""].value_counts()
        if not vc.empty:
            course_mapping["{{C_CLASS}}"] = vc.idxmax()
        else:
            course_mapping["{{C_CLASS}}"] = ""
    else:
        course_mapping["{{C_CLASS}}"] = ""

    all_columns = list(df.columns)
    if len(all_columns) <= 4:
        print("沒有評分項目欄位（第 5 欄之後）。")
        return

    # 假設前四欄：學號, 姓名, 科系, 班級
    item_columns = all_columns[4:]

    os.makedirs(output_dir, exist_ok=True)
    random.seed()

    item_serial = 0  # 評量序號：0,1,2,...

    for col in item_columns:
        col_name = str(col)
        scores_raw = student_df[col].tolist()

        if not should_select_column(col_name, scores_raw):
            continue

        parsed_scores = [parse_score(v) for v in scores_raw]

        expected = num_students
        submitted = sum(1 for s in parsed_scores if s is not None)
        missing = expected - submitted

        valid_students = []
        for (_, row), s in zip(student_df.iterrows(), parsed_scores):
            if s is None:
                continue
            valid_students.append({
                "id": row["學號"],
                "name": row["姓名"],
                "score": float(s),
            })

        if not valid_students:
            print(f"[略過] 項目「{col_name}」沒有任何有效成績。")
            continue

        scores_only = [stu["score"] for stu in valid_students]
        full_avg = sum(scores_only) / len(scores_only)

        segA, segB, segC = split_into_three_segments(valid_students)
        segments_stats = {}
        for label, seg in zip(["A", "B", "C"], [segA, segB, segC]):
            if seg:
                seg_scores = [s["score"] for s in seg]
                segments_stats[label] = {
                    "students": seg,
                    "count": len(seg),
                    "max": max(seg_scores),
                    "min": min(seg_scores),
                    "avg": sum(seg_scores) / len(seg_scores),
                }
            else:
                segments_stats[label] = {
                    "students": [],
                    "count": 0,
                    "max": 0.0,
                    "min": 0.0,
                    "avg": 0.0,
                }

        stats = {
            "expected": expected,
            "submitted": submitted,
            "missing": missing,
            "full_avg": full_avg,
            "segments": segments_stats,
        }

        # 每個欄位專屬：I_ID / I_TYPE
        item_meta = parse_item_meta(col_name)
        item_meta["I_ID"] = str(item_serial)
        item_serial += 1

        mapping = build_mapping_for_item(col_name, stats, course_mapping, item_meta)

        # 檔名安全化
        base_name = "".join(c if c not in r'\/:*?"<>|' else "_" for c in col_name)
        output_path = os.path.join(output_dir, f"{base_name}_繳交記錄表.docx")

        fill_template(template_path, output_path, mapping, item_meta)

        print("=" * 80)
        print(f"已產生：{output_path}")
        print(f"  項目名稱：{col_name}")
        print(f"  評量類別(I_TYPE)：{item_meta.get('I_TYPE', '')}")
        print(f"  作業序號(I_ID)：{item_meta.get('I_ID', '')}")
        print(f"  應交：{expected}  實交(有效成績)：{submitted}  缺交：{missing}")
        print(f"  全班有效成績平均：{full_avg:.2f}")
        print(f"  課程名稱：{course_mapping.get('{{C_NAME}}', '')}")
        print(f"  授課教師：{course_mapping.get('{{C_TEACHER}}', '')}")
        print(f"  班級(多數)：{course_mapping.get('{{C_CLASS}}', '')}")


# ================== 主程式 (argparse) ================== #

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="IEET Score Table Generator with Course Info")
    parser.add_argument(
        "excel_path",
        type=str,
        help="成績 Excel 檔案路徑，例如 IEET_score_list_108520_3.xlsx"
    )
    parser.add_argument(
        "--template",
        type=str,
        default="IEET_TEMPLATE_86_112X.docx",
        help="Word 範本 (含 placeholder)，預設為 IEET_TEMPLATE_86_112X.docx"
    )
    parser.add_argument(
        "--sheet",
        type=str,
        default=None,
        help="指定工作表名稱；若不指定則使用第一個工作表"
    )
    args = parser.parse_args()

    excel_path = args.excel_path
    template_path = args.template
    sheet_name = args.sheet

    # output_dir = "IEET_OUTPUT" + excel 檔名（去副檔名）
    base = os.path.basename(excel_path)
    name_no_ext = os.path.splitext(base)[0]
    output_dir = f"IEET_OUTPUT_{name_no_ext}"

    os.makedirs(output_dir, exist_ok=True)

    analyze_and_generate(
        excel_path=excel_path,
        template_path=template_path,
        output_dir=output_dir,
        sheet_name=sheet_name
    )

    print("\n=== 全部完成 ===")
    print(f"輸出資料夾：{output_dir}")
